"""
runner.py — The single unified run loop for one Task on one provider.

Architecture:
    1. Start MCPToolClient (which launches exec_server.py via stdio)
    2. Discover tools from the MCP server
    3. Run agentic loop: model calls → tool routing via MCP → results back
    4. Capture full trajectory (every turn, every tool call, every result)

Files are NOT inlined in the prompt. The model discovers and reads
them via MCP tools (python_execute, bash_execute, etc.).
"""

from __future__ import annotations

import os
import logging
from datetime import datetime, timezone

try:
    from . import file_gen
    from .provider import supports_tools
    from .models import RunResult, ToolCall, TurnRecord
    from .mcp_client import MCPToolClient
except ImportError:
    import file_gen
    from provider import supports_tools
    from models import RunResult, ToolCall, TurnRecord
    from mcp_client import MCPToolClient

logger = logging.getLogger("dra.runner")

SYSTEM_PROMPT = (
    "You are a deep research analyst. Produce a comprehensive, well-structured, "
    "well-cited report based on the user's question and any provided files. "
    "Cite non-trivial claims with numbered references. "
    "Do not fabricate sources.\n\n"
    "You have access to tools for reading files, executing code, running shell commands, "
    "searching the web, and writing output files. Use them as needed. "
    "All input files are in your current working directory. Write every output file "
    "to the current working directory using a plain filename — do not use absolute "
    "paths and do not 'cd' elsewhere, or your files will be lost between steps."
)


# ─── Message construction ─────────────────────────────────────────────────────

def build_messages(task, params) -> list[dict]:
    """Build initial messages. Model discovers files via tools.

    Parity note: the only file-related instruction is a single static line
    naming the expected deliverable format(s) — this mirrors how APEX/GDPVal
    task definitions state the required deliverable. No sentinel blocks, no
    library hand-holding, no reactive coaching.
    """
    parts = [task.prompt]

    if task.file_paths:
        parts.append(
            "\n\nReference files have been provided in your working directory. "
            "Use your tools to list, read, and analyze them."
        )

    # One static deliverable instruction, built from the task's expected format.
    # Filenames use plain names in the working directory (see system prompt).
    if params.file_output and task.output_formats:
        fmts = ", ".join(f".{f}" for f in task.output_formats)
        parts.append(
            f"\n\nProduce your deliverable(s) as {fmts} file(s) written to your "
            f"working directory. The deliverable must contain the finished output "
            f"the task asks for — not your reasoning, notes, or code."
        )

    return [
        {"role": "system", "content": SYSTEM_PROMPT},
        {"role": "user", "content": "\n".join(parts)},
    ]

# def build_messages(task, params) -> list[dict]:
#     """
#     Build the initial [system, user] messages.

#     Files are NOT inlined. The model discovers and reads them via tools.
#     """
#     parts = [task.prompt]

#     if task.file_paths:
#         parts.append("\n" + "=" * 60)
#         parts.append("AVAILABLE REFERENCE FILES (use tools to read them)")
#         parts.append("=" * 60)
#         for fpath in task.file_paths:
#             basename = os.path.basename(fpath)
#             size = os.path.getsize(fpath) if os.path.exists(fpath) else 0
#             parts.append(f"  - {basename}  ({size:,} bytes)")
#         parts.append(
#             "\nUse python_execute or bash_execute to read and analyze these files. "
#             "They are in the current working directory."
#         )
#         parts.append("=" * 60)

#     if params.file_output and task.output_formats:
#         parts.append(file_gen.file_gen_instructions(task.output_formats[0], task.file_paths))

#     return [
#         {"role": "system", "content": SYSTEM_PROMPT},
#         {"role": "user", "content": "\n".join(parts)},
#     ]


# ─── Stage input files ───────────────────────────────────────────────────────

def _stage_input_files(file_paths: list, run_dir: str):
    """Symlink input files into the run directory. Fast, no copy."""
    os.makedirs(run_dir, exist_ok=True)
    for src in file_paths:
        if not os.path.exists(src):
            continue
        dst = os.path.join(run_dir, os.path.basename(src))
        if os.path.exists(dst):
            continue
        try:
            os.symlink(os.path.abspath(src), dst)
        except OSError:
            # Windows or filesystem doesn't support symlinks — fall back to copy
            import shutil
            shutil.copy2(src, dst)


# ─── Harvest model-generated deliverables ─────────────────────────────────────

# Deliverable file types the model may write via write_file / python_execute.
_DELIVERABLE_EXTS = {".md", ".txt", ".json", ".csv", ".docx", ".xlsx", ".pptx", ".pdf"}
# Transient helper scripts the model writes to do its work — not deliverables.
_TRANSIENT_NAMES = {"calc.py", "make_memo.py"}


def _harvest_output_files(staging_dir: str, input_paths: list) -> list[str]:
    """Collect genuine model-generated files from the staging dir.

    Input files are symlinks (see _stage_input_files); model outputs are real
    files. We skip symlinks, input basenames (even if the model overwrote the
    symlink with a real file of the same name — those are NOT deliverables), and
    transient .py scripts. Runs regardless of file_output / forced_stop, so a
    capped-but-productive run still reports what it managed to write.
    """
    input_names = {os.path.basename(p) for p in (input_paths or [])}
    found = []
    for root, _dirs, files in os.walk(staging_dir):
        for name in sorted(files):
            full = os.path.join(root, name)
            if os.path.islink(full):            # input symlink
                continue
            if name in input_names or name in _TRANSIENT_NAMES:
                # Same basename as an input → treat as input, not a deliverable.
                # (If a task legitimately needs to emit a same-named file, have
                #  the prompt require an 'output_' prefix; see build_messages.)
                continue
            if name.startswith("tmp") and name.endswith(".py"):
                continue
            if os.path.splitext(name)[1].lower() not in _DELIVERABLE_EXTS:
                continue
            found.append(full)
    return found


# ─── Content validation (pure harness verdict — no model call, no feedback) ───
#
# APEX-parity note: this ONLY inspects and classifies. It never re-prompts the
# model or triggers regeneration. A file judged to be reasoning/code is recorded
# as an invalid deliverable — the same verdict a downstream scorer would reach —
# and the task simply scores as "no valid deliverable." That is the honest APEX
# outcome, not a second chance.

# Line-leading markers that signal reasoning/narration rather than a deliverable.
_REASONING_MARKERS = (
    "let me", "i'll", "i will", "first,", "next,", "now i", "let's",
    "i need to", "i should", "okay,", "ok,", "step 1", "step 2",
    "here's my", "here is my", "to solve this", "my analysis",
    "thinking", "reasoning:",
)
# Line-leading markers that signal raw code.
_CODE_MARKERS = (
    "import ", "from ", "def ", "class ", "print(", "```", "#!/",
    "if __name__", "return ", "for ", "while ", "try:", "except",
)


def _extract_text_for_validation(path: str) -> str | None:
    """Best-effort text pull from a deliverable for content inspection.
    Returns None if the type can't be inspected as text (treated as OK)."""
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext in (".txt", ".md", ".csv", ".json"):
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                return f.read()
        if ext == ".docx":
            from docx import Document
            doc = Document(path)
            return "\n".join(p.text for p in doc.paragraphs)
        # xlsx/pptx/pdf: structural formats — a reasoning dump is far less
        # likely and harder to false-positive on, so we don't text-scan them.
        return None
    except Exception:
        return None


def _validate_deliverable_content(path: str, response_text: str) -> tuple[bool, str]:
    """Return (is_valid, reason). Pure judgment — does NOT modify or regenerate.

    Flags a file as invalid when it looks like reasoning/narration or raw code
    rather than a finished deliverable, or when it's essentially a copy of the
    model's reasoning trace (response_text).
    """
    text = _extract_text_for_validation(path)
    if text is None:
        return True, "not text-inspectable; accepted"
    stripped = text.strip()
    if len(stripped) < 50:
        return False, "too short to be a deliverable"

    lines = [ln.strip() for ln in stripped.splitlines() if ln.strip()]
    if not lines:
        return False, "no content lines"

    reasoning_hits = sum(
        1 for ln in lines if ln.lower().startswith(_REASONING_MARKERS)
    )
    code_hits = sum(
        1 for ln in lines if ln.lstrip().startswith(_CODE_MARKERS)
    )
    frac_reasoning = reasoning_hits / len(lines)
    frac_code = code_hits / len(lines)

    if frac_code > 0.30:
        return False, f"looks like code ({frac_code:.0%} code-like lines)"
    if frac_reasoning > 0.30:
        return False, f"looks like reasoning ({frac_reasoning:.0%} narration lines)"

    # Near-duplicate of the reasoning trace → it's the transcript, not a deliverable.
    if response_text and len(response_text) > 200:
        rt = response_text.strip()
        # cheap similarity: how much of the file is the leading reasoning text
        if stripped[:500] and stripped[:500] in rt:
            return False, "content mirrors the model's reasoning trace"

    return True, "ok"


# ─── Deliverable post-processing (md→docx, response→docx, model prefix) ───────

def _md_to_docx(md_path: str) -> str | None:
    """Convert a Markdown file to .docx alongside it, return the new path.
    Best-effort: headings (#..), bullet/numbered lists, and paragraphs. Falls
    back to plain paragraphs if structure isn't recognized. Returns None on error.
    """
    try:
        from docx import Document
    except ImportError:
        logger.warning("python-docx not installed — cannot convert %s", md_path)
        return None
    try:
        with open(md_path, "r", encoding="utf-8") as f:
            lines = f.read().splitlines()
        doc = Document()
        for raw in lines:
            line = raw.rstrip()
            if not line:
                continue
            stripped = line.lstrip()
            if stripped.startswith("#"):
                level = len(stripped) - len(stripped.lstrip("#"))
                text = stripped[level:].strip()
                doc.add_heading(text, level=min(level, 4))
            elif stripped[:2] in ("- ", "* ", "+ "):
                doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            elif len(stripped) > 2 and stripped[0].isdigit() and stripped[1:3] in (". ", ") "):
                doc.add_paragraph(stripped[3:].strip(), style="List Number")
            else:
                # strip inline ** bold / * markers lightly for readability
                doc.add_paragraph(stripped.replace("**", "").replace("`", ""))
        out = os.path.splitext(md_path)[0] + ".docx"
        doc.save(out)
        return out
    except Exception as e:
        logger.warning("md→docx failed for %s: %s", md_path, e)
        return None


def _response_to_docx(response_text: str, out_path: str) -> str | None:
    """Render the model's response_text as a readable .docx. Treats the text as
    Markdown-ish (same heuristics as _md_to_docx)."""
    if not response_text or not response_text.strip():
        return None
    try:
        from docx import Document
    except ImportError:
        return None
    try:
        doc = Document()
        for raw in response_text.splitlines():
            stripped = raw.strip()
            if not stripped:
                continue
            if stripped.startswith("#"):
                level = len(stripped) - len(stripped.lstrip("#"))
                doc.add_heading(stripped[level:].strip(), level=min(level, 4))
            elif stripped[:2] in ("- ", "* ", "+ "):
                doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            else:
                doc.add_paragraph(stripped.replace("**", "").replace("`", ""))
        doc.save(out_path)
        return out_path
    except Exception as e:
        logger.warning("response→docx failed: %s", e)
        return None


def _postprocess_outputs(harvested: list, staging: str, task, response_text: str) -> list[str]:
    """Apply the output policy to harvested files:
      1. Convert every .md deliverable to .docx and DROP the .md.
      2. Always emit {provider}_{task_id}_response.docx from response_text.
      3. Prefix every output filename with the provider (so multi-model runs
         don't collide), unless already prefixed.
    Returns the final list of output file paths.
    """
    final: list[str] = []
    prefix = f"{task.provider}_"

    # 1. md → docx (drop the .md)
    converted = []
    for p in harvested:
        if p.lower().endswith(".md"):
            docx_path = _md_to_docx(p)
            if docx_path:
                try:
                    os.remove(p)           # drop the .md per policy
                except OSError:
                    pass
                converted.append(docx_path)
            else:
                converted.append(p)        # conversion failed; keep .md rather than lose it
        else:
            converted.append(p)

    # 2. response_text → its own docx (always, separate from deliverables)
    resp_name = f"{prefix}{task.task_id}_response.docx"
    resp_path = os.path.join(staging, resp_name)
    made = _response_to_docx(response_text, resp_path)
    if made:
        converted.append(made)

    # 3. model-prefix every filename (skip if already prefixed)
    for p in converted:
        d, name = os.path.split(p)
        if name.startswith(prefix):
            final.append(p)
            continue
        new_path = os.path.join(d, prefix + name)
        try:
            os.rename(p, new_path)
            final.append(new_path)
        except OSError:
            final.append(p)                # rename failed; keep original path
    return final


# ─── Run one task ─────────────────────────────────────────────────────────────

async def run_task(task, params, driver) -> RunResult:
    started = datetime.now(timezone.utc)
    logger.info(
        "[%s] start (model=%s, dry_run=%s)",
        task.run_id, task.model_slug, driver.dry_run,
    )

    if driver.dry_run:
        return _dry_run_result(task, params, started)

    # ── Set up staging directory ──────────────────────────────────
    staging = task.output_dir
    if not staging:
        raise RuntimeError(f"task {task.task_id} has no output_dir — pipeline.build_tasks must set it")
    os.makedirs(staging, exist_ok=True)
    _stage_input_files(task.file_paths, staging)

    messages = build_messages(task, params)

    # ── Start MCP client ──────────────────────────────────────────
    mcp: MCPToolClient | None = None
    schemas = None

    if supports_tools(task.provider):
        try:
            mcp = MCPToolClient(staging_dir=staging)
            await mcp.start()
            schemas = mcp.openai_schemas
        except Exception as e:
            logger.warning(
                "[%s] MCP failed to start: %s — running without tools",
                task.run_id, e,
            )
            mcp = None
            schemas = None

    # ── Agentic loop ──────────────────────────────────────────────
    in_tok = out_tok = 0
    cost = 0.0
    turns = 0
    trajectory: list[TurnRecord] = []
    tool_log: list[ToolCall] = []
    citations: list = []
    seen_urls: set = set()
    final_text = ""
    forced_stop = False
    error = None

    def _add_citations(cits):
        for c in cits or []:
            url = c.get("url", "")
            if url and url not in seen_urls:
                citations.append(c)
                seen_urls.add(url)

    try:
        for turn in range(1, params.max_turns + 1):
            turns = turn
            turn_ts = datetime.now(timezone.utc).isoformat()

            resp = await driver.chat(task.model_slug, messages, params, schemas)

            in_tok += resp.input_tokens
            out_tok += resp.output_tokens
            cost += resp.cost_usd
            _add_citations(resp.citations)
            final_text = resp.text or final_text

            # ── Build turn record ─────────────────────────────────
            record = TurnRecord(
                turn=turn,
                timestamp=turn_ts,
                assistant_text=resp.text or "",
                finish_reason=resp.finish_reason,
                input_tokens=resp.input_tokens,
                output_tokens=resp.output_tokens,
                cost_usd=resp.cost_usd,
            )

            # Budget guard
            if cost > params.max_cost_usd:
                logger.warning(
                    "[%s] budget $%.2f exceeded ($%.2f) — stopping",
                    task.run_id, params.max_cost_usd, cost,
                )
                forced_stop = True
                trajectory.append(record)
                break

            # No tool calls → final answer
            if not resp.tool_calls or not schemas:
                trajectory.append(record)
                break

            # ── Execute tool calls via MCP ────────────────────────
            messages.append(resp.assistant_message)

            for tc in resp.tool_calls:
                tc_name = tc["name"]
                tc_args = tc["arguments"]
                tc_id = tc["id"]

                # Record the call
                record.tool_calls.append({
                    "name": tc_name,
                    "arguments": tc_args,
                    "id": tc_id,
                })

                # Route through MCP client
                if mcp and mcp.is_ready:
                    result = await mcp.call_tool(tc_name, tc_args)
                else:
                    result = f"[error] MCP not available for: {tc_name}"

                # Record the result
                record.tool_results.append({
                    "name": tc_name,
                    "result": result[:10000],
                    "server": "exec",
                })

                # Legacy flat log
                tool_log.append(ToolCall(
                    turn=turn,
                    name=tc_name,
                    arguments=tc_args,
                    result_preview=result[:500],
                ))

                # Feed back to model
                messages.append({
                    "role": "tool",
                    "tool_call_id": tc_id,
                    "name": tc_name,
                    "content": result,
                })

            trajectory.append(record)

        else:
            forced_stop = True
            logger.warning("[%s] max_turns (%d) reached", task.run_id, params.max_turns)

    except Exception as e:
        logger.exception("[%s] run error", task.run_id)
        error = str(e)
        forced_stop = True

    finally:
        if mcp:
            await mcp.stop()

    # ── File generation ───────────────────────────────────────────
    output_files: list = []
    output_file_errors: dict = {}
    if params.file_output and task.output_formats and final_text and not forced_stop:
        output_files, output_file_errors, fix_cost = await file_gen.generate(
            final_text, task, params, driver, task.model_slug
        )
        cost += fix_cost

    # ── Harvest + post-process what the model wrote via MCP tools ─────
    # Independent of file_gen and of forced_stop: a run that hit the turn cap
    # after writing a deliverable should still report that file. Deduped against
    # anything file_gen already produced.
    harvested = _harvest_output_files(staging, task.file_paths)
    for f in harvested:
        if f not in output_files:
            output_files.append(f)

    # Apply output policy: md→docx (drop md), always add response docx,
    # prefix every filename with the provider.
    output_files = _postprocess_outputs(
        output_files, staging, task, final_text
    )

    # ── Content validation (pure verdict; no regeneration, no model feedback) ──
    # Judge each real deliverable: does it read like a finished deliverable, or
    # like reasoning/code? Invalid files are recorded and excluded from the
    # "real deliverable" set, so a task that only produced junk scores as
    # "no valid deliverable" — the honest APEX outcome, with no second chance.
    resp_docx = f"{task.provider}_{task.task_id}_response.docx"
    deliverable_validity: dict = {}   # basename -> {"valid": bool, "reason": str}
    for f in output_files:
        base = os.path.basename(f)
        if base == resp_docx:
            continue   # convenience artifact, not a measured deliverable
        is_valid, reason = _validate_deliverable_content(f, final_text)
        deliverable_validity[base] = {"valid": is_valid, "reason": reason}
        if not is_valid:
            logger.warning("[%s] deliverable '%s' rejected: %s",
                           task.run_id, base, reason)

    # Completion: a run is complete if it produced text AND at least one VALID
    # deliverable (when the task demanded a file). The response.docx never counts
    # toward this. Invalid files stay in output_files (for inspection) but don't
    # satisfy the requirement.
    demands_file = bool(task.output_formats)
    valid_deliverables = [
        f for f in output_files
        if os.path.basename(f) != resp_docx
        and deliverable_validity.get(os.path.basename(f), {}).get("valid", True)
    ]
    produced_file = bool(valid_deliverables)
    completed = (
        bool(final_text)
        and not output_file_errors
        and (produced_file or not demands_file)
        and (not forced_stop or produced_file)
    )
    if output_file_errors and not error:
        error = f"File generation failed: {list(output_file_errors.keys())}"
    # Surface content rejections in the error field when they cause incompletion.
    rejected = [b for b, v in deliverable_validity.items() if not v["valid"]]
    if demands_file and not produced_file and rejected and not error:
        error = f"deliverable content rejected: {rejected}"

    return _finish(
        task, started,
        response_text=final_text,
        citations=citations,
        input_tokens=in_tok,
        output_tokens=out_tok,
        cost=cost,
        turns=turns,
        tool_log=tool_log,
        trajectory=trajectory,
        completed=completed,
        forced_stop=forced_stop,
        error=error,
        output_files=output_files,
        output_file_errors=output_file_errors,
    )


# ─── Result builders ──────────────────────────────────────────────────────────

def _finish(task, started, *, response_text, citations, input_tokens, output_tokens,
            cost, turns, tool_log, trajectory, completed, forced_stop, error,
            output_files=None, output_file_errors=None) -> RunResult:
    completed_at = datetime.now(timezone.utc)
    return RunResult(
        task_id=task.task_id,
        run_id=task.run_id,
        provider=task.provider,
        model=task.model_slug,
        pass_index=task.pass_index,
        response_text=response_text,
        citations=citations,
        output_files=output_files or [],
        output_file_errors=output_file_errors or {},
        tool_calls=[tc.to_dict() for tc in tool_log],
        trajectory=[t.to_dict() for t in trajectory],
        input_tokens=input_tokens,
        output_tokens=output_tokens,
        total_cost_usd=round(cost, 6),
        turns=turns,
        completed=completed,
        forced_stop=forced_stop,
        error=error,
        total_duration_sec=(completed_at - started).total_seconds(),
        started_at=started.isoformat(),
        completed_at=completed_at.isoformat(),
    )


def _dry_run_result(task, params, started) -> RunResult:
    fmts = ", ".join(task.output_formats) if task.output_formats else "none"
    files = ", ".join(os.path.basename(f) for f in task.file_paths) or "none"
    report = (
        f"# [DRY RUN] {task.provider} via OpenRouter\n\n"
        f"**Task:** {task.task_id}  \n**Model:** {task.model_slug}  \n"
        f"**Pass:** {task.pass_index}  \n**Web search:** {params.web_search}  \n"
        f"**Tools:** MCP (exec_server)  \n"
        f"**Output formats:** {fmts}  \n**Files:** {files}  \n"
        f"**max_turns:** {params.max_turns}  \n\n"
        f"## Prompt\n\n{task.prompt[:500]}{'...' if len(task.prompt) > 500 else ''}\n"
    )
    return _finish(
        task, started,
        response_text=report,
        citations=[],
        input_tokens=0,
        output_tokens=0,
        cost=0.0,
        turns=0,
        tool_log=[],
        trajectory=[],
        completed=True,
        forced_stop=False,
        error=None,
        output_files=[],
        output_file_errors=({task.output_formats[0]: "DRY_RUN"}
                           if task.output_formats else {}),
    )