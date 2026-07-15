"""
mcp_server.py — MCP server exposing all DRA tools.

Started as a subprocess by the runner. Communicates via stdio transport.
Tools are discovered dynamically by the MCP client — no hardcoded schemas
in the agent code, matching APEX-Agents architecture.

    python mcp_server.py               # starts the stdio server
    python mcp_server.py --list-tools  # print available tools and exit
"""

import os
import sys
import json
import subprocess
import tempfile
import logging

from fastmcp import FastMCP

logger = logging.getLogger("dra.mcp_server")

mcp = FastMCP(
    name="dra-tools",
    instructions=(
        "Tools for deep research analysis. Use python_execute for calculations, "
        "read_file for loading data, web_search for external information, "
        "and write_file for producing deliverables."
    ),
)

STAGING_DIR = os.environ.get("INDRAYUDH_STAGING_DIR", os.getcwd())


# ─── Python execution ─────────────────────────────────────────────────────────

@mcp.tool()
def python_execute(code: str) -> str:
    """Execute Python code and return stdout/stderr. Use for calculations,
    data analysis, file processing, chart generation. Common libraries
    available: pandas, numpy, scipy, openpyxl, matplotlib, json, csv.
    The working directory contains the task input files."""

    with tempfile.NamedTemporaryFile(
        mode="w", suffix=".py", dir=STAGING_DIR, delete=False
    ) as f:
        f.write(code)
        tmp_path = f.name

    try:
        result = subprocess.run(
            [sys.executable, tmp_path],
            capture_output=True, text=True, timeout=120, cwd=STAGING_DIR,
        )
        output = ""
        if result.stdout:
            output += result.stdout
        if result.stderr:
            output += f"\n[stderr]\n{result.stderr}"
        return (output.strip() or "[no output]")[:15000]
    except subprocess.TimeoutExpired:
        return "[error] execution timed out (120s)"
    except Exception as e:
        return f"[error] {e}"
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


# ─── Shell execution ──────────────────────────────────────────────────────────

@mcp.tool()
def bash_execute(command: str) -> str:
    """Execute a shell command and return stdout/stderr. Use for file
    operations, installing packages, running scripts, or system commands."""

    try:
        result = subprocess.run(
            command, shell=True,
            capture_output=True, text=True, timeout=120, cwd=STAGING_DIR,
        )
        output = ""
        if result.stdout:
            output += result.stdout
        if result.stderr:
            output += f"\n[stderr]\n{result.stderr}"
        return (output.strip() or "[no output]")[:15000]
    except subprocess.TimeoutExpired:
        return "[error] execution timed out (120s)"
    except Exception as e:
        return f"[error] {e}"


# ─── Read file ────────────────────────────────────────────────────────────────

@mcp.tool()
def read_file(path: str) -> str:
    """Read file contents. Supports .txt, .csv, .json, .md, .py, .xlsx
    (all sheets), .pdf (text extraction), .docx. Returns text truncated
    to 50,000 characters."""

    if not os.path.isabs(path):
        path = os.path.join(STAGING_DIR, path)

    if not os.path.isfile(path):
        return f"[error] file not found: {path}"

    ext = os.path.splitext(path)[1].lower()
    max_chars = 50000

    try:
        if ext in (".txt", ".csv", ".json", ".md", ".py", ".js", ".html",
                   ".xml", ".yaml", ".yml", ".toml", ".log", ".tsv", ".sql"):
            with open(path, "r", encoding="utf-8", errors="replace") as f:
                return f.read()[:max_chars]

        if ext in (".xlsx", ".xls"):
            import openpyxl
            wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
            parts = []
            for sn in wb.sheetnames:
                ws = wb[sn]
                parts.append(f"=== Sheet: {sn} ===")
                for row in ws.iter_rows(values_only=True):
                    parts.append("\t".join(
                        str(c) if c is not None else "" for c in row
                    ))
            wb.close()
            return "\n".join(parts)[:max_chars]

        if ext == ".pdf":
            import pdfplumber
            text_parts = []
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        text_parts.append(t)
            return "\n\n".join(text_parts)[:max_chars] or "[no extractable text]"

        if ext == ".docx":
            from docx import Document
            doc = Document(path)
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    text += "\n" + "\t".join(c.text.strip() for c in row.cells)
            return text[:max_chars] or "[empty document]"

        with open(path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()[:max_chars]

    except Exception as e:
        return f"[error] reading {path}: {e}"


# ─── Search in file ───────────────────────────────────────────────────────────

@mcp.tool()
def search_in_file(path: str, query: str, context_lines: int = 3) -> str:
    """Search for a keyword or phrase within a file and return matching
    lines with surrounding context. Use when a file is large and you need
    to find specific data without reading the entire file."""

    content = read_file(path)
    if content.startswith("[error]"):
        return content

    lines = content.split("\n")
    query_lower = query.lower()
    matches = []

    for i, line in enumerate(lines):
        if query_lower in line.lower():
            start = max(0, i - context_lines)
            end = min(len(lines), i + context_lines + 1)
            snippet = "\n".join(
                f"{'>>>' if j == i else '   '} {lines[j]}"
                for j in range(start, end)
            )
            matches.append(f"[Match at line {i+1}]\n{snippet}")

    if not matches:
        return f"No matches for '{query}' in {os.path.basename(path)}"

    header = f"Found {len(matches)} match(es) for '{query}':\n\n"
    return (header + "\n\n".join(matches))[:15000]


# ─── Write file ───────────────────────────────────────────────────────────────

@mcp.tool()
def write_file(path: str, content: str) -> str:
    """Write text content to a file. Use for creating reports, saving
    analysis results, generating CSV/JSON output, or writing code files."""

    if not os.path.isabs(path):
        path = os.path.join(STAGING_DIR, path)

    try:
        os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)
        return f"File written: {path} ({os.path.getsize(path):,} bytes)"
    except Exception as e:
        return f"[error] writing {path}: {e}"


# ─── List directory ───────────────────────────────────────────────────────────

@mcp.tool()
def list_directory(path: str = ".") -> str:
    """List files and subdirectories. Defaults to the task staging directory."""

    if not os.path.isabs(path):
        path = os.path.join(STAGING_DIR, path)

    if not os.path.isdir(path):
        return f"[error] not a directory: {path}"

    try:
        entries = []
        for name in sorted(os.listdir(path)):
            full = os.path.join(path, name)
            if os.path.isdir(full):
                entries.append(f"  {name}/")
            else:
                entries.append(f"  {name}  ({os.path.getsize(full):,} bytes)")
        return f"Directory: {path}\n" + "\n".join(entries) if entries else "[empty]"
    except Exception as e:
        return f"[error] listing {path}: {e}"


# ─── Web search ───────────────────────────────────────────────────────────────

@mcp.tool()
def web_search(query: str, max_results: int = 5) -> str:
    """Search the web using DuckDuckGo. Returns titles, URLs, and snippets.
    Use for finding current data, verifying facts, or retrieving external
    information not in the provided files."""

    try:
        from duckduckgo_search import DDGS
    except ImportError:
        return "[error] pip install duckduckgo-search"

    try:
        with DDGS() as ddgs:
            results = list(ddgs.text(query, max_results=max_results))

        if not results:
            return f"No results for '{query}'"

        parts = []
        for i, r in enumerate(results, 1):
            parts.append(
                f"[{i}] {r.get('title', '')}\n"
                f"    URL: {r.get('href', '')}\n"
                f"    {r.get('body', '')}"
            )
        return "\n\n".join(parts)
    except Exception as e:
        return f"[error] search failed: {e}"


# ─── Web fetch ────────────────────────────────────────────────────────────────

@mcp.tool()
def web_fetch(url: str) -> str:
    """Fetch text content of a web page. Use for retrieving data from APIs,
    reading web pages, or downloading publicly available documents.
    Returns content truncated to 30,000 characters."""

    import urllib.request
    import urllib.error
    import re

    try:
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        with urllib.request.urlopen(req, timeout=30) as resp:
            content_type = resp.headers.get("Content-Type", "")
            raw = resp.read(500_000)
            encoding = "utf-8"
            if "charset=" in content_type:
                encoding = content_type.split("charset=")[-1].split(";")[0].strip()
            text = raw.decode(encoding, errors="replace")

            if "html" in content_type.lower():
                text = re.sub(r"<script[^>]*>.*?</script>", "", text, flags=re.DOTALL)
                text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.DOTALL)
                text = re.sub(r"<[^>]+>", " ", text)
                text = re.sub(r"\s+", " ", text).strip()

            return text[:30000] or "[empty response]"

    except Exception as e:
        return f"[error] fetching {url}: {e}"


# ─── Calculator ───────────────────────────────────────────────────────────────

@mcp.tool()
def calculator(expression: str) -> str:
    """Evaluate a basic arithmetic expression and return the result."""

    import ast
    import operator as op

    ops = {
        ast.Add: op.add, ast.Sub: op.sub, ast.Mult: op.mul,
        ast.Div: op.truediv, ast.Pow: op.pow, ast.Mod: op.mod,
        ast.USub: op.neg, ast.UAdd: op.pos, ast.FloorDiv: op.floordiv,
    }

    def _eval(node):
        if isinstance(node, ast.Constant) and isinstance(node.value, (int, float)):
            return node.value
        if isinstance(node, ast.BinOp) and type(node.op) in ops:
            return ops[type(node.op)](_eval(node.left), _eval(node.right))
        if isinstance(node, ast.UnaryOp) and type(node.op) in ops:
            return ops[type(node.op)](_eval(node.operand))
        raise ValueError("unsupported expression")

    try:
        tree = ast.parse(expression, mode="eval")
        return str(_eval(tree.body))
    except Exception as e:
        return f"[error] {e}"


# ─── Entry point ──────────────────────────────────────────────────────────────

if __name__ == "__main__":
    if "--list-tools" in sys.argv:
        for tool in mcp._tool_manager.tools.values():
            print(f"  {tool.name}: {tool.description[:80]}")
        sys.exit(0)

    mcp.run(transport="stdio")